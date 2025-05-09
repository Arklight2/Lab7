# firm/views.py
import re
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse
from django.core.mail import send_mail
from django.conf import settings
from django.contrib.auth.decorators import login_required
import openpyxl
from docx import Document
from reportlab.pdfgen import canvas
from .forms import RegistrationForm, ClientForm
from .models import CustomUser, Client
from .utils import custom_hash


def register(request):
    if request.method == "POST":
        form = RegistrationForm(request.POST)
        if form.is_valid():
            # Создаём пользователя и хешируем пароль с помощью custom_hash
            user = form.save(commit=False)
            raw_password = form.cleaned_data.get('password1')
            user.password = custom_hash(raw_password, algorithm="sha256")
            user.save()
            messages.success(request, "Регистрация успешна!")
            return redirect('login')
        else:
            messages.error(request, "Исправьте ошибки в форме.")
    else:
        form = RegistrationForm()
    return render(request, "firm/register.html", {"form": form})


def password_recovery(request):
    if request.method == "POST":
        email = request.POST.get('email')
        try:
            user = CustomUser.objects.get(email=email)
            # Генерируем токен на основе email и случайной строки (salt)
            token = custom_hash(email + "secret_salt", algorithm="sha256")
            recovery_link = request.build_absolute_uri(f"/reset_password/{token}/")
            send_mail(
                "Восстановление пароля",
                f"Перейдите по ссылке для восстановления пароля: {recovery_link}",
                settings.DEFAULT_FROM_EMAIL,
                [email],
            )
            messages.success(request, "Инструкции по восстановлению пароля отправлены на e-mail!")
        except CustomUser.DoesNotExist:
            messages.error(request, "Пользователь с таким e-mail не найден.")
    return render(request, "firm/password_recovery.html")


@login_required
def client_list(request):
    # Обычный пользователь видит только свои записи
    if request.user.role == 'user':
        clients = Client.objects.filter(created_by=request.user)
    else:
        clients = Client.objects.all()
    return render(request, "firm/client_list.html", {"clients": clients})


@login_required
def add_client(request):
    if request.method == "POST":
        form = ClientForm(request.POST)
        if form.is_valid():
            client = form.save(commit=False)
            client.created_by = request.user
            client.save()
            messages.success(request, "Клиент успешно добавлен!")
            return redirect('client_list')
        else:
            messages.error(request, "Исправьте ошибки в форме.")
    else:
        form = ClientForm()
    return render(request, "firm/client_form.html", {"form": form})


@login_required
def update_client(request, pk):
    client = get_object_or_404(Client, pk=pk, created_by=request.user)
    if request.method == "POST":
        form = ClientForm(request.POST, instance=client)
        if form.is_valid():
            form.save()
            messages.success(request, "Запись обновлена успешно!")
            return redirect('client_list')
        else:
            messages.error(request, "Исправьте ошибки в форме.")
    else:
        form = ClientForm(instance=client)
    return render(request, "firm/client_update.html", {"form": form})


@login_required
def delete_client(request, pk):
    client = get_object_or_404(Client, pk=pk, created_by=request.user)
    if request.method == "POST":
        client.delete()
        messages.success(request, "Запись удалена успешно!")
        return redirect('client_list')
    return render(request, "firm/client_confirm_delete.html", {"client": client})


@login_required
def export_clients_excel(request):
    clients = Client.objects.all() if request.user.role == 'admin' else Client.objects.filter(created_by=request.user)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Клиенты"
    ws.append(["ID", "Фамилия", "Имя", "E-mail", "Дата регистрации"])

    for client in clients:
        ws.append([client.id, client.surname, client.name, client.email,
                   client.registration_date.strftime("%d.%m.%Y %H:%M")])

    response = HttpResponse(content_type="application/vnd.ms-excel")
    response['Content-Disposition'] = 'attachment; filename="clients.xlsx"'
    wb.save(response)
    return response


@login_required
def export_clients_word(request):
    clients = Client.objects.all() if request.user.role == 'admin' else Client.objects.filter(created_by=request.user)
    document = Document()
    document.add_heading('Список клиентов', 0)

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Фамилия"
    hdr_cells[2].text = "Имя"
    hdr_cells[3].text = "E-mail"
    hdr_cells[4].text = "Дата регистрации"

    for client in clients:
        row_cells = table.add_row().cells
        row_cells[0].text = str(client.id)
        row_cells[1].text = client.surname
        row_cells[2].text = client.name
        row_cells[3].text = client.email
        row_cells[4].text = client.registration_date.strftime("%d.%m.%Y %H:%M")

    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = 'attachment; filename="clients.docx"'
    document.save(response)
    return response


from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from reportlab.pdfgen import canvas
from .models import Client


@login_required
def export_clients_pdf(request):
    # Получаем список клиентов: для администратора — все, для обычного — только свои
    if request.user.role == 'admin':
        clients = Client.objects.all()
    else:
        clients = Client.objects.filter(created_by=request.user)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="clients.pdf"'

    p = canvas.Canvas(response)

    # Начинаем с первой страницы: задаём заголовок
    p.setFont("Helvetica", 14)
    p.drawString(100, 800, "Список клиентов")

    y = 770
    p.setFont("Helvetica", 10)
    for client in clients:
        # Формируем строку с данными клиента
        client_line = f"{client.id} | {client.surname} {client.name} | {client.email} | {client.registration_date.strftime('%d.%m.%Y %H:%M')}"
        p.drawString(50, y, client_line)
        y -= 20
        # Если места на странице мало, начинаем новую страницу
        if y < 50:
            p.showPage()
            p.setFont("Helvetica", 14)
            p.drawString(100, 800, "Список клиентов (продолжение)")
            y = 770
            p.setFont("Helvetica", 10)

    p.save()
    return response


